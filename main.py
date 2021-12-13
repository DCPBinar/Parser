import os
import zipfile
import xml.dom.minidom
import csv
import re
import docx

cols = ['Имя', 'Фамилия', 'Полугодие', 'Оценка 1', 'Оценка 2', 'Оценка 3', 'Оценка 4', 'Оценка 5', 'Оценка 6', 'Оценка 7']

n = 0

for root, dirs, files in os.walk('C:/tables/'):
    for file in files:
        n += 1
        os.rename(f'C:/tables/{file}', f'C:/tables/{n}.docx')

files = os.listdir(path='C:/tables/')


names = []

for n in range(len(files)):
    doc = docx.Document(f'C:/tables/{n + 1}.docx')
    document = zipfile.ZipFile(f'C:/tables/{n + 1}.docx')
    XML = xml.dom.minidom.parseString(document.read('word/document.xml')).toprettyxml(indent='  ')

    text_re = re.compile('>\n\s+([^<>\s].*?)\n\s+</', re.DOTALL)
    prettyXml = text_re.sub('>\g<1></', XML)

    file = open(f'C:/tables/{n + 1}.xml', 'w+', encoding='utf-8')
    file.write(prettyXml)
    file.close()

    first_name = re.findall('<w:t xml:space="preserve">.+</w:t>', prettyXml)
    first_name = str(first_name[0][27:-6])
    last_name = re.findall('<w:t>.+</w:t>', prettyXml)
    last_name = str(last_name[0][5:-6])

    first_table = doc.tables[5]
    second_table = doc.tables[2]

    a = []
    for row in first_table.rows:
        string = ''
        for cell in row.cells:
            string = string + cell.text + ' '
        a.append(string)

    b = []
    for row in second_table.rows:
        string = ''
        for cell in row.cells:
            string = string + cell.text + ' '
        b.append(string)

    data = [[first_name, last_name, '1', str(a[1][-2]), str(a[2][-2]), str(a[3][-2]), str(a[4][-2]), str(a[5][-2]), str(a[6][-2]), str(a[7][-2])],
            [first_name, last_name, '2', str(b[1][-2]), str(b[2][-2]), str(b[3][-2]), str(b[4][-2]), str(b[5][-2]), str(b[6][-2]), str(b[7][-2])]]

    with open('C:/tables/output.csv', 'w', newline='') as f:
        writer = csv.writer(f, delimiter=';')
        writer.writerows(data)

